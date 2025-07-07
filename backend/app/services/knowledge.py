"""
Knowledge Service
Handles knowledge base management, document processing, and search
"""

import logging
from typing import List, Dict, Optional, Any
from datetime import datetime
import asyncio
import hashlib
import json

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_
import openai
from PyPDF2 import PdfReader
from docx import Document
import io

from app.core.config import settings
from app.models.knowledge_base import KnowledgeBase, KnowledgeDocument
from app.models.tenant import Tenant


logger = logging.getLogger(__name__)


class KnowledgeService:
    """Service for managing knowledge bases and document search"""
    
    def __init__(self):
        self.openai_client = openai.AsyncOpenAI(api_key=settings.OPENAI_API_KEY)
    
    async def create_knowledge_base(
        self, 
        tenant_id: str, 
        name: str, 
        description: str,
        db: AsyncSession
    ) -> KnowledgeBase:
        """Create a new knowledge base for a tenant"""
        
        knowledge_base = KnowledgeBase(
            tenant_id=tenant_id,
            name=name,
            description=description,
            is_active=True,
            is_default=False
        )
        
        db.add(knowledge_base)
        await db.commit()
        await db.refresh(knowledge_base)
        
        logger.info(f"Created knowledge base '{name}' for tenant {tenant_id}")
        return knowledge_base
    
    async def upload_document(
        self, 
        tenant_id: str, 
        knowledge_base_id: str,
        file_content: bytes, 
        filename: str, 
        content_type: str,
        db: AsyncSession
    ) -> KnowledgeDocument:
        """Upload and process a document to the knowledge base"""
        
        # Verify knowledge base belongs to tenant
        stmt = select(KnowledgeBase).where(
            and_(
                KnowledgeBase.id == knowledge_base_id,
                KnowledgeBase.tenant_id == tenant_id
            )
        )
        result = await db.execute(stmt)
        kb = result.scalar_one_or_none()
        
        if not kb:
            raise ValueError("Knowledge base not found or access denied")
        
        # Extract text content
        text_content = await self._extract_text_content(
            file_content, filename, content_type
        )
        
        if not text_content:
            raise ValueError("Could not extract text from document")
        
        # Generate content hash for deduplication
        content_hash = hashlib.sha256(text_content.encode()).hexdigest()
        
        # Check if document already exists
        stmt = select(KnowledgeDocument).where(
            and_(
                KnowledgeDocument.knowledge_base_id == knowledge_base_id,
                KnowledgeDocument.content_hash == content_hash
            )
        )
        result = await db.execute(stmt)
        existing_doc = result.scalar_one_or_none()
        
        if existing_doc:
            logger.info(f"Document with same content already exists: {filename}")
            return existing_doc
        
        # Create document record
        document = KnowledgeDocument(
            knowledge_base_id=knowledge_base_id,
            filename=filename,
            content_type=content_type,
            file_size=len(file_content),
            content_hash=content_hash,
            raw_content=text_content,
            status="processing"
        )
        
        db.add(document)
        await db.commit()
        await db.refresh(document)
        
        # Process document in background
        asyncio.create_task(
            self._process_document_chunks(document.id, text_content, db)
        )
        
        logger.info(f"Uploaded document '{filename}' to knowledge base {knowledge_base_id}")
        return document
    
    async def search_knowledge(
        self, 
        tenant_id: str, 
        query: str, 
        db: AsyncSession,
        knowledge_base_id: Optional[str] = None,
        limit: int = 5
    ) -> List[Dict[str, Any]]:
        """Search knowledge base using semantic similarity"""
        
        try:
            # Get query embedding
            query_embedding = await self._get_text_embedding(query)
            
            # Build base query
            stmt = select(KnowledgeDocument).join(KnowledgeBase).where(
                and_(
                    KnowledgeBase.tenant_id == tenant_id,
                    KnowledgeBase.is_active == True,
                    KnowledgeDocument.status == "completed"
                )
            )
            
            if knowledge_base_id:
                stmt = stmt.where(KnowledgeDocument.knowledge_base_id == knowledge_base_id)
            
            result = await db.execute(stmt)
            documents = result.scalars().all()
            
            if not documents:
                return []
            
            # Calculate similarity scores (simplified - in production use vector DB)
            results = []
            for doc in documents:
                if doc.processed_chunks:
                    chunks = json.loads(doc.processed_chunks)
                    for chunk in chunks:
                        # Simple text similarity (in production, use embeddings)
                        similarity_score = self._calculate_text_similarity(
                            query.lower(), chunk["content"].lower()
                        )
                        
                        if similarity_score > 0.1:  # Minimum relevance threshold
                            results.append({
                                "document_id": str(doc.id),
                                "filename": doc.filename,
                                "content": chunk["content"],
                                "similarity_score": similarity_score,
                                "chunk_index": chunk["index"]
                            })
            
            # Sort by similarity score and return top results
            results.sort(key=lambda x: x["similarity_score"], reverse=True)
            return results[:limit]
            
        except Exception as e:
            logger.error(f"Knowledge search error: {e}")
            return []
    
    async def get_knowledge_bases(
        self, 
        tenant_id: str, 
        db: AsyncSession
    ) -> List[KnowledgeBase]:
        """Get all knowledge bases for a tenant"""
        
        stmt = select(KnowledgeBase).where(
            and_(
                KnowledgeBase.tenant_id == tenant_id,
                KnowledgeBase.is_active == True
            )
        ).order_by(KnowledgeBase.created_at.desc())
        
        result = await db.execute(stmt)
        return result.scalars().all()
    
    async def get_documents(
        self, 
        knowledge_base_id: str, 
        tenant_id: str,
        db: AsyncSession
    ) -> List[KnowledgeDocument]:
        """Get all documents in a knowledge base"""
        
        stmt = select(KnowledgeDocument).join(KnowledgeBase).where(
            and_(
                KnowledgeDocument.knowledge_base_id == knowledge_base_id,
                KnowledgeBase.tenant_id == tenant_id
            )
        ).order_by(KnowledgeDocument.created_at.desc())
        
        result = await db.execute(stmt)
        return result.scalars().all()
    
    async def delete_document(
        self, 
        document_id: str, 
        tenant_id: str,
        db: AsyncSession
    ) -> bool:
        """Delete a document from knowledge base"""
        
        stmt = select(KnowledgeDocument).join(KnowledgeBase).where(
            and_(
                KnowledgeDocument.id == document_id,
                KnowledgeBase.tenant_id == tenant_id
            )
        )
        result = await db.execute(stmt)
        document = result.scalar_one_or_none()
        
        if not document:
            return False
        
        await db.delete(document)
        await db.commit()
        
        logger.info(f"Deleted document {document_id}")
        return True
    
    async def _extract_text_content(
        self, 
        file_content: bytes, 
        filename: str, 
        content_type: str
    ) -> str:
        """Extract text content from various file types"""
        
        try:
            if content_type == "application/pdf" or filename.endswith(".pdf"):
                return await self._extract_pdf_text(file_content)
            
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ] or filename.endswith(".docx"):
                return await self._extract_docx_text(file_content)
            
            elif content_type == "text/plain" or filename.endswith(".txt"):
                return file_content.decode("utf-8")
            
            elif filename.endswith(".md"):
                return file_content.decode("utf-8")
            
            else:
                # Try to decode as text
                return file_content.decode("utf-8")
                
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            return ""
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF text extraction error: {e}")
            return ""
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX text extraction error: {e}")
            return ""
    
    async def _get_text_embedding(self, text: str) -> List[float]:
        """Get text embedding using OpenAI"""
        try:
            response = await self.openai_client.embeddings.create(
                model="text-embedding-ada-002",
                input=text
            )
            return response.data[0].embedding
            
        except Exception as e:
            logger.error(f"Embedding generation error: {e}")
            return []
    
    async def _process_document_chunks(
        self, 
        document_id: str, 
        content: str, 
        db: AsyncSession
    ):
        """Process document into searchable chunks"""
        try:
            # Split content into chunks
            chunks = self._split_text_into_chunks(content, chunk_size=1000, overlap=100)
            
            # Process each chunk
            processed_chunks = []
            for i, chunk in enumerate(chunks):
                # In production, you'd generate embeddings for each chunk
                processed_chunks.append({
                    "index": i,
                    "content": chunk,
                    "word_count": len(chunk.split()),
                    "char_count": len(chunk)
                })
            
            # Update document with processed chunks
            stmt = select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
            result = await db.execute(stmt)
            document = result.scalar_one_or_none()
            
            if document:
                document.processed_chunks = json.dumps(processed_chunks)
                document.status = "completed"
                document.processed_at = datetime.utcnow()
                await db.commit()
            
            logger.info(f"Processed document {document_id} into {len(chunks)} chunks")
            
        except Exception as e:
            logger.error(f"Document processing error: {e}")
            # Mark document as failed
            stmt = select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
            result = await db.execute(stmt)
            document = result.scalar_one_or_none()
            if document:
                document.status = "failed"
                document.error_message = str(e)
                await db.commit()
    
    def _split_text_into_chunks(
        self, 
        text: str, 
        chunk_size: int = 1000, 
        overlap: int = 100
    ) -> List[str]:
        """Split text into overlapping chunks"""
        
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = " ".join(chunk_words)
            chunks.append(chunk_text)
            
            if i + chunk_size >= len(words):
                break
        
        return chunks
    
    def _calculate_text_similarity(self, query: str, content: str) -> float:
        """Calculate simple text similarity (for demo - use embeddings in production)"""
        
        query_words = set(query.split())
        content_words = set(content.split())
        
        if not query_words or not content_words:
            return 0.0
        
        intersection = query_words.intersection(content_words)
        union = query_words.union(content_words)
        
        return len(intersection) / len(union) if union else 0.0


# Global knowledge service instance
knowledge_service = KnowledgeService()